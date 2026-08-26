import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE

doc = docx.Document()

# Set standard margins (0.75 in all sides)
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Exact Original Executive Palette
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79)    # #1F4E79 Deep Executive Blue
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # #2E75B6 Medium Accent Blue
COLOR_MUTED_LINE = RGBColor(0xAA, 0xAA, 0xAA) # #AAAAAA Soft Muted Grey
COLOR_BODY = RGBColor(0x59, 0x59, 0x59)       # #595959 Slate Grey Body
COLOR_LINK = "0563C1"                         # #0563C1 Hyperlink Blue

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}"/>')
    new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    rPr = parse_xml(f'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    # Blue color + single underline
    c = parse_xml(f'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="{COLOR_LINK}"/>')
    u = parse_xml(f'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>')
    rPr.append(c)
    rPr.append(u)
    
    # Font settings
    sz = parse_xml(f'<w:sz xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="19"/>') # 9.5pt
    rPr.append(sz)
    
    new_run.append(rPr)
    
    # Add text node
    t = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve">{text}</w:t>')
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# Header Paragraph
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(2)
title_run = title_p.add_run("ALEX TEYE AMETEPEY")
title_run.font.name = 'Calibri'
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(6)
sub_run = sub_p.add_run("Full-Stack Software Engineer & Systems Architect")
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(12)
sub_run.font.bold = True
sub_run.font.color.rgb = COLOR_ACCENT

# Contact Bar with Clickable Hyperlinks and Muted Separators
contact_p = doc.add_paragraph()
contact_p.paragraph_format.space_before = Pt(0)
contact_p.paragraph_format.space_after = Pt(14)

def add_plain_contact(text):
    r = contact_p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_BODY

def add_sep():
    r = contact_p.add_run("  |  ")
    r.font.name = 'Calibri'
    r.font.size = Pt(9.5)
    r.font.color.rgb = COLOR_MUTED_LINE

add_plain_contact("Accra, Ghana")
add_sep()
add_plain_contact("0549044977")
add_sep()
add_hyperlink(contact_p, "mailto:alexteyeametepey@gmail.com", "alexteyeametepey@gmail.com")
add_sep()
add_hyperlink(contact_p, "https://github.com/aa-Teye", "github.com/aa-Teye")
add_sep()
add_hyperlink(contact_p, "https://www.linkedin.com/in/alex-ametepey-1123a3205", "linkedin.com/in/alex-ametepey-1123a3205")
add_sep()
add_hyperlink(contact_p, "https://alex-portfolio-sooty.vercel.app", "alex-portfolio-sooty.vercel.app")

# Add Header Bottom Border Line (#1F4E79)
header_pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="8" w:space="5" w:color="1F4E79"/></w:pBdr>')
contact_p._p.get_or_add_pPr().append(header_pBdr)

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    
    # Add bottom border accent line (#2E75B6) under section header
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="3" w:color="2E75B6"/></w:pBdr>')
    p._p.get_or_add_pPr().append(pBdr)
    return p

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25)
    
    b_run = p.add_run("• ")
    b_run.font.name = 'Calibri'
    b_run.font.size = Pt(9.5)
    b_run.font.color.rgb = COLOR_PRIMARY
    
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_PRIMARY
    
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = COLOR_BODY
    return p

# 1. Professional Summary
add_heading(doc, "Professional Summary")
sum_p = doc.add_paragraph()
sum_p.paragraph_format.space_before = Pt(3)
sum_p.paragraph_format.space_after = Pt(8)
sum_p.paragraph_format.line_spacing = 1.15
sum_run = sum_p.add_run(
    "Results-driven Full-Stack Software Engineer and Systems Architect with a strong track record of leading technical "
    "teams to deliver end-to-end enterprise platforms, mobile applications, and automated growth systems. Transitioning "
    "a background in machine learning and backend architecture into comprehensive full-stack development to solve complex "
    "operational challenges. Currently pursuing an MSc in Financial Engineering to integrate quantitative data modeling "
    "into production-grade software architectures."
)
sum_run.font.name = 'Calibri'
sum_run.font.size = Pt(9.5)
sum_run.font.color.rgb = COLOR_BODY

# 2. Technical Skills
add_heading(doc, "Technical Skills")
add_bullet(doc, "Languages: ", "Python, JavaScript (ES6+), Java, etc.")
add_bullet(doc, "Frontend & Mobile: ", "React, Next.js, Astro, Vite, Tailwind CSS, Cross-Platform Mobile Development (React Native), Electron")
add_bullet(doc, "Backend & Architecture: ", "FastAPI, RESTful API Design, System Design, JWT Authentication, Redis (Caching & Queues)")
add_bullet(doc, "Databases & Storage: ", "Neon (Serverless PostgreSQL), PostgreSQL, SQLite, Relational Schema Design")
add_bullet(doc, "DevOps & Workflows: ", "Vercel, Render, Git, GitHub, Jira, Agile Sprint Planning")
add_bullet(doc, "Media Systems: ", "Live Broadcast Routing (vMix), OBS, Freeshow, etc.")

# 3. Key Projects
add_heading(doc, "Key Full-Stack & System Architecture Projects")
add_bullet(doc, "Ghana HIVdr Platform (Ongoing): ", "Lead the full-stack architecture and data synchronization pipelines for secure, real-time health-tech reporting and record tracking.")
add_bullet(doc, "UGDS Customer Care System: ", "Led the full-stack architecture, API endpoint specifications, and database performance optimization for real-time customer care and ticket escalation.")
add_bullet(doc, "Learning and Examination Management System: ", "Built a secure assessment engine supporting automated grading logic, audit logging, and encrypted academic record handling.")
add_bullet(doc, "EcoPulse Web Platform [Win, mobile and web app, and God's Eye IoT integration]: ", "Developed an ecological analytics and sustainability tracking platform displaying real-time environmental metrics and data visualizations.")
add_bullet(doc, "Enterprise ERP & Custom CMS Solutions: ", "Architected multi-tenant database schemas, automated operational workflows, and role-based access control (RBAC) systems across diverse client requirements.")
add_bullet(doc, "Church Management & Operations System (SPS): ", "Engineered an integrated platform for database administration, member tracking, and automated service coordination.")

# 4. Professional Experience
add_heading(doc, "Professional Experience")

def add_role_header(doc, role, company, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    
    # Right-aligned tab stop at 7.0 in for dates
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    
    r_role = p.add_run(role)
    r_role.font.name = 'Calibri'
    r_role.font.bold = True
    r_role.font.size = Pt(10.5)
    r_role.font.color.rgb = COLOR_PRIMARY
    
    r_sep = p.add_run("  |  ")
    r_sep.font.name = 'Calibri'
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = COLOR_MUTED_LINE
    
    r_comp = p.add_run(company)
    r_comp.font.name = 'Calibri'
    r_comp.font.bold = True
    r_comp.font.size = Pt(10)
    r_comp.font.color.rgb = COLOR_ACCENT
    
    r_date = p.add_run(f"\t{dates}")
    r_date.font.name = 'Calibri'
    r_date.font.size = Pt(9.5)
    r_date.font.color.rgb = COLOR_BODY

add_role_header(doc, "Software Developer (Contract)", "Aréte Forge", "Ongoing")
add_bullet(doc, "", "Architect and deploy comprehensive full-stack solutions, engineering seamless experiences from responsive web platforms to mobile applications.")
add_bullet(doc, "", "Maintain modular backend microservices using Python and FastAPI, bridging APIs with modern frontend interfaces.")

add_role_header(doc, "Backend Developer (Contract)", "Meditel", "Ongoing")
add_bullet(doc, "", "Design and deploy secure RESTful endpoints and high-performance backend services using Python and FastAPI.")
add_bullet(doc, "", "Collaborate closely with frontend engineers to integrate backend services seamlessly into React-based applications.")

add_role_header(doc, "Growth Engineer (Contract)", "WAICA", "Ongoing")
add_bullet(doc, "", "Design and implement automated customer acquisition funnels, technical marketing systems, and digital growth campaigns.")
add_bullet(doc, "", "Leverage data analytics and automation tools to optimize conversion rates and expand digital brand presence.")

add_role_header(doc, "Head of IT & Media Systems", "Overcomers Nation Church (ONC)", "2023 - Present")
add_bullet(doc, "", "Direct IT operations, advanced AV media broadcasting systems, and internal technical infrastructure.")
add_bullet(doc, "", "Develop custom full-stack web and mobile applications to automate administrative workflows, community engagement, and digital production.")

add_role_header(doc, "Research Assistant", "University of Ghana", "2023 - Present")
add_bullet(doc, "", "Conduct deep learning research under Prof. Kofi Sarpong Adu-Manu focusing on spatio-temporal graph transformer networks.")
add_bullet(doc, "", "Build automated data pipelines and evaluate model training benchmarks to support advanced AI infrastructure monitoring.")

# 5. Education
add_heading(doc, "Education")

def add_edu_item(doc, degree, institution, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    
    r_deg = p.add_run(degree)
    r_deg.font.name = 'Calibri'
    r_deg.font.bold = True
    r_deg.font.size = Pt(10.5)
    r_deg.font.color.rgb = COLOR_PRIMARY
    
    r_sep = p.add_run("  |  ")
    r_sep.font.name = 'Calibri'
    r_sep.font.color.rgb = COLOR_MUTED_LINE
    
    r_inst = p.add_run(institution)
    r_inst.font.name = 'Calibri'
    r_inst.font.size = Pt(10)
    r_inst.font.color.rgb = COLOR_BODY
    
    r_date = p.add_run(f"\t{dates}")
    r_date.font.name = 'Calibri'
    r_date.font.size = Pt(9.5)
    r_date.font.color.rgb = COLOR_BODY

add_edu_item(doc, "MSc in Financial Engineering", "WorldQuant University (USA)", "Ongoing")
add_edu_item(doc, "BSc in Information Technology", "University of Ghana, Legon", "Graduated 2025")

# 6. References
add_heading(doc, "References")

def add_ref_item(doc, name, title_org, contact):
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(1)
    p1.paragraph_format.keep_with_next = True
    
    r_name = p1.add_run(name)
    r_name.font.name = 'Calibri'
    r_name.font.bold = True
    r_name.font.size = Pt(10)
    r_name.font.color.rgb = COLOR_PRIMARY
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    
    r_info = p2.add_run(f"{title_org}  |  {contact}")
    r_info.font.name = 'Calibri'
    r_info.font.size = Pt(9)
    r_info.font.color.rgb = COLOR_BODY

add_ref_item(doc, "Professor Kofi Sarpong Adu-Manu", "Dept. of Computer Science, University of Ghana", "kaysarpsnr@gmail.com / 0244602374")
add_ref_item(doc, "Alex Quao, PhD", "CEO, Aréte Forge", "0592199757 / 0249221772")
add_ref_item(doc, "Dr. Ebenezer Okronipa", "Pharmacist, General Overseer of EOM", "0200994446")

# Save DOCX files
output_path1 = "Alex_Teye_Ametepey_CV.docx"
output_path2 = "public/Alex_Resume.docx"
doc.save(output_path1)
doc.save(output_path2)

print(f"Successfully generated {output_path1} and {output_path2} with exact right-aligned tab stops and executive formatting!")
